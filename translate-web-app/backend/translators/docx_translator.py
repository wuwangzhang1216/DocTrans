"""
Word (DOCX) translation module.
Preserves formatting, styles, and table structures.
"""

import os
import json
from typing import List
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
from google import genai
from pydantic import BaseModel
from docx import Document

from .core import TranslationClient


class TranslatedRuns(BaseModel):
    """Schema for translated text runs"""
    translations: List[str]


class DOCXTranslator:
    """
    Word document translator that preserves formatting and styles.
    """

    def __init__(self, translation_client: TranslationClient):
        """
        Initialize DOCX translator.

        Args:
            translation_client: Core translation client instance
        """
        self.client = translation_client

    def _translate_runs(self, runs: List[str], target_language: str,
                       api_key: str, model: str, max_retries: int = 3) -> List[str]:
        """
        Translate runs while preserving structure, with retry on failure.
        Uses Pydantic schema for guaranteed JSON format.

        Args:
            runs: List of text runs
            target_language: Target language
            api_key: API key
            model: Model name
            max_retries: Maximum number of retry attempts

        Returns:
            List of translated runs
        """
        if not runs:
            return runs

        if api_key:
            os.environ['GEMINI_API_KEY'] = api_key
        genai_client = genai.Client()

        prompt = (
            f"You are a professional technical translator. Translate the following text runs into {target_language}.\n\n"
            "Rules:\n"
            f"- Translate each run into {target_language} with native, accurate, technical wording\n"
            "- Consider all runs together for context, but return the translation per run\n"
            "- Keep the EXACT same number of runs in the same order\n"
            "- Do not merge or split runs; preserve whitespace\n"
            "- Preserve placeholders, code, variables, URLs, and IDs unchanged\n"
            "- Return exactly the same number of translations as input runs\n\n"
            f"Input runs ({len(runs)} items):\n{json.dumps(runs, ensure_ascii=False)}"
        )

        # Retry logic for API failures
        import time
        for attempt in range(max_retries):
            try:
                resp = genai_client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config={
                        "response_mime_type": "application/json",
                        "response_schema": TranslatedRuns,
                    }
                )

                # Use parsed response for guaranteed structure
                result: TranslatedRuns = resp.parsed

                if len(result.translations) == len(runs):
                    return result.translations
                else:
                    print(f"Warning: Translation count mismatch. Expected {len(runs)}, got {len(result.translations)}")
                    # Retry if count doesn't match
                    if attempt < max_retries - 1:
                        continue

            except Exception as e:
                if attempt < max_retries - 1:
                    # Exponential backoff: wait 1s, 2s, 4s...
                    wait_time = 2 ** attempt
                    print(f"Translation attempt {attempt + 1} failed: {str(e)}. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    print(f"Translation failed after {max_retries} attempts: {str(e)}")

        return runs

    def _translate_single_paragraph(self, para_data: tuple) -> tuple:
        """
        Translate a single paragraph (to be called in parallel).

        Args:
            para_data: Tuple containing (para_idx, runs, target_language, api_key, model)

        Returns:
            Tuple of (para_idx, translated_runs)
        """
        para_idx, runs, target_language, api_key, model = para_data

        try:
            if runs and any(t.strip() for t in runs):
                translated_runs = self._translate_runs(runs, target_language, api_key, model)
            else:
                translated_runs = runs
            return (para_idx, [(i, t) for i, t in enumerate(translated_runs)])
        except Exception as e:
            print(f"Translation error in paragraph {para_idx}: {str(e)}")
            return (para_idx, [(i, t) for i, t in enumerate(runs)])

    def _translate_paragraph_batch(self, batch_data: tuple) -> dict:
        """
        Translate a batch of paragraphs in parallel.

        Args:
            batch_data: Tuple containing (batch_index, paragraphs_runs_data, target_language, api_key, model, max_workers)

        Returns:
            Dictionary with translated paragraphs
        """
        batch_idx, paragraphs_runs_data, target_language, api_key, model, batch_max_workers = batch_data

        try:
            results = []

            # Prepare data for parallel processing
            para_tasks = [
                (para_idx, runs, target_language, api_key, model)
                for para_idx, runs in paragraphs_runs_data
            ]

            # Process paragraphs in parallel within this batch
            with ThreadPoolExecutor(max_workers=batch_max_workers) as executor:
                results = list(executor.map(self._translate_single_paragraph, para_tasks))

            return {"batch_idx": batch_idx, "translated_paragraphs": results}

        except Exception as e:
            print(f"Translation error in batch {batch_idx}: {str(e)}")
            return {"batch_idx": batch_idx, "error": str(e)}

    def _translate_single_table_paragraph(self, table_para_data: tuple) -> tuple:
        """
        Translate a single table paragraph (to be called in parallel).

        Args:
            table_para_data: Tuple containing (key, runs, target_language, api_key, model)

        Returns:
            Tuple of (key, translated_runs)
        """
        key, runs, target_language, api_key, model = table_para_data

        try:
            if runs and any(t.strip() for t in runs):
                translated_runs = self._translate_runs(runs, target_language, api_key, model)
            else:
                translated_runs = runs
            return (key, [(i, t) for i, t in enumerate(translated_runs)])
        except Exception as e:
            print(f"Translation error in table paragraph {key}: {str(e)}")
            return (key, [(i, t) for i, t in enumerate(runs)])

    def _translate_table_batch(self, batch_data: tuple) -> dict:
        """
        Translate a batch of table paragraphs in parallel.

        Args:
            batch_data: Tuple containing (batch_index, table_para_runs_data, target_language, api_key, model, max_workers)

        Returns:
            Dictionary with translated table paragraphs
        """
        batch_idx, table_para_runs_data, target_language, api_key, model, batch_max_workers = batch_data

        try:
            results = []

            # Prepare data for parallel processing
            table_para_tasks = [
                (key, runs, target_language, api_key, model)
                for key, runs in table_para_runs_data
            ]

            # Process table paragraphs in parallel within this batch
            with ThreadPoolExecutor(max_workers=batch_max_workers) as executor:
                results = list(executor.map(self._translate_single_table_paragraph, table_para_tasks))

            return {"batch_idx": batch_idx, "translated_table_paragraphs": results}

        except Exception as e:
            print(f"Translation error in table batch {batch_idx}: {str(e)}")
            return {"batch_idx": batch_idx, "error": str(e)}

    def translate(self, input_path: str, output_path: str,
                 target_language: str) -> bool:
        """
        Translate Word document.

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save translated DOCX file
            target_language: Target language for translation

        Returns:
            Success status
        """
        try:
            doc = Document(input_path)

            # Gather paragraph runs
            body_paragraphs = []
            for para_idx, paragraph in enumerate(doc.paragraphs):
                runs_texts = [r.text for r in paragraph.runs]
                body_paragraphs.append((para_idx, runs_texts))

            # Gather table cell paragraph runs
            table_para_runs = []
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            runs_texts = [r.text for r in paragraph.runs]
                            table_para_runs.append(((table_idx, row_idx, cell_idx, para_idx), runs_texts))

            # Calculate worker allocation for maximum parallelism
            total_items = len(body_paragraphs) + len(table_para_runs)
            max_workers = self.client.max_workers  # 256
            max_concurrent_batches = self.client.max_concurrent_pages  # 16

            # Strategy: Create max_concurrent_batches batches, each with its own worker pool
            # Total workers = max_concurrent_batches * workers_per_batch ≤ max_workers
            target_batches = min(max_concurrent_batches, total_items)
            workers_per_batch = max(1, max_workers // target_batches) if target_batches > 0 else max_workers
            items_per_batch = max(1, total_items // target_batches) if target_batches > 0 else total_items

            print(f"Processing DOCX with maximum parallelism:")
            print(f"  - Body paragraphs: {len(body_paragraphs)}")
            print(f"  - Table paragraphs: {len(table_para_runs)}")
            print(f"  - Concurrent batches: {target_batches}")
            print(f"  - Items per batch: ~{items_per_batch}")
            print(f"  - Workers per batch: {workers_per_batch}")
            print(f"  - Total concurrent API calls: {target_batches * workers_per_batch}")

            # Process body paragraphs in batches
            paragraph_batches = []
            if body_paragraphs:
                # Use API key from the translation client
                api_key = self.client.api_key if hasattr(self.client, 'api_key') and self.client.api_key else os.environ.get('GEMINI_API_KEY', None)
                # Create batches of paragraphs
                for i in range(0, len(body_paragraphs), items_per_batch):
                    batch = body_paragraphs[i:i + items_per_batch]
                    paragraph_batches.append((len(paragraph_batches), batch, target_language, api_key, self.client.model, workers_per_batch))

            translated_body = {}
            if paragraph_batches:
                print(f"Processing {len(paragraph_batches)} paragraph batches...")
                # Use target_batches for outer parallelism (batch-level)
                with ThreadPoolExecutor(max_workers=target_batches) as executor:
                    future_to_batch = {
                        executor.submit(self._translate_paragraph_batch, batch): batch[0]
                        for batch in paragraph_batches
                    }

                    for future in concurrent.futures.as_completed(future_to_batch):
                        batch_idx = future_to_batch[future]
                        try:
                            result = future.result()
                            if 'error' not in result:
                                for para_idx, translated_runs in result['translated_paragraphs']:
                                    translated_body[para_idx] = translated_runs
                        except Exception as e:
                            print(f"Paragraph batch {batch_idx} translation failed: {str(e)}")

            # Process table paragraphs in batches
            table_batches = []
            if table_para_runs:
                # Use API key from the translation client
                api_key = self.client.api_key if hasattr(self.client, 'api_key') and self.client.api_key else os.environ.get('GEMINI_API_KEY', None)
                # Create batches of table paragraphs
                for i in range(0, len(table_para_runs), items_per_batch):
                    batch = table_para_runs[i:i + items_per_batch]
                    table_batches.append((len(table_batches), batch, target_language, api_key, self.client.model, workers_per_batch))

            translated_table = {}
            if table_batches:
                print(f"Processing {len(table_batches)} table batches...")
                # Use target_batches for outer parallelism (batch-level)
                with ThreadPoolExecutor(max_workers=target_batches) as executor:
                    future_to_batch = {
                        executor.submit(self._translate_table_batch, batch): batch[0]
                        for batch in table_batches
                    }

                    for future in concurrent.futures.as_completed(future_to_batch):
                        batch_idx = future_to_batch[future]
                        try:
                            result = future.result()
                            if 'error' not in result:
                                for key, translated_runs in result['translated_table_paragraphs']:
                                    translated_table[key] = translated_runs
                        except Exception as e:
                            print(f"Table batch {batch_idx} translation failed: {str(e)}")

            # Apply body paragraph translations
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if para_idx in translated_body:
                    for run_idx, run_text in translated_body[para_idx]:
                        if run_idx < len(paragraph.runs):
                            paragraph.runs[run_idx].text = run_text

            # Apply table translations
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            key = (table_idx, row_idx, cell_idx, para_idx)
                            if key in translated_table:
                                for run_idx, run_text in translated_table[key]:
                                    if run_idx < len(paragraph.runs):
                                        paragraph.runs[run_idx].text = run_text

            # Save translated document
            doc.save(output_path)
            print(f"[SUCCESS] Translated DOCX saved to: {output_path}")
            return True

        except Exception as e:
            print(f"Error translating DOCX: {str(e)}")
            return False
