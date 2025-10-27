"""
Word (DOCX) translation module.
Preserves formatting, styles, and table structures.

Performance Optimizations:
- Single-level parallelization with all available workers (256)
- Combined processing of body and table paragraphs (no sequential bottleneck)
- Adaptive to document size - near 100% worker utilization
- Expected speedup: 10-50x for typical documents (10-500 paragraphs)
"""

import os
import json
from typing import List
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures
from google import genai
from pydantic import BaseModel
from docx import Document

from .core import TranslationClient


class TranslatedRuns(BaseModel):
    """Schema for translated text runs"""
    translations: List[str]


class DOCXTranslator:
    """
    Word document translator that preserves formatting and styles.
    """

    def __init__(self, translation_client: TranslationClient):
        """
        Initialize DOCX translator.

        Args:
            translation_client: Core translation client instance
        """
        self.client = translation_client

    def _translate_runs(self, runs: List[str], target_language: str,
                       api_key: str, model: str, max_retries: int = 3) -> List[str]:
        """
        Translate runs while preserving structure, with retry on failure.
        Uses Pydantic schema for guaranteed JSON format.

        Args:
            runs: List of text runs
            target_language: Target language
            api_key: API key
            model: Model name
            max_retries: Maximum number of retry attempts

        Returns:
            List of translated runs
        """
        if not runs:
            return runs

        # Ensure API key is set in environment for this thread
        if api_key:
            os.environ['GEMINI_API_KEY'] = api_key

        # Create client with explicit API key
        genai_client = genai.Client(api_key=api_key if api_key else os.environ.get('GEMINI_API_KEY'))

        prompt = (
            f"You are a professional technical translator. Translate the following text runs into {target_language}.\n\n"
            "Rules:\n"
            f"- Translate each run into {target_language} with native, accurate, technical wording\n"
            "- Consider all runs together for context, but return the translation per run\n"
            "- Keep the EXACT same number of runs in the same order\n"
            "- Do not merge or split runs; preserve whitespace\n"
            "- Preserve placeholders, code, variables, URLs, and IDs unchanged\n"
            "- Return exactly the same number of translations as input runs\n\n"
            f"Input runs ({len(runs)} items):\n{json.dumps(runs, ensure_ascii=False)}"
        )

        # Retry logic for API failures
        import time
        for attempt in range(max_retries):
            try:
                resp = genai_client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config={
                        "response_mime_type": "application/json",
                        "response_schema": TranslatedRuns,
                    }
                )

                # Use parsed response for guaranteed structure
                result: TranslatedRuns = resp.parsed

                if len(result.translations) == len(runs):
                    print(f"[DEBUG] Translation successful: {len(runs)} runs translated")
                    return result.translations
                else:
                    print(f"Warning: Translation count mismatch. Expected {len(runs)}, got {len(result.translations)}")
                    # Retry if count doesn't match
                    if attempt < max_retries - 1:
                        continue

            except Exception as e:
                if attempt < max_retries - 1:
                    # Exponential backoff: wait 1s, 2s, 4s...
                    wait_time = 2 ** attempt
                    print(f"Translation attempt {attempt + 1} failed: {str(e)}. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    print(f"Translation failed after {max_retries} attempts: {str(e)}")

        print(f"[DEBUG] Returning original runs after all attempts failed")
        return runs

    def _translate_single_paragraph(self, para_data: tuple) -> tuple:
        """
        Translate a single paragraph (to be called in parallel).

        Args:
            para_data: Tuple containing (para_idx, runs, target_language, api_key, model)

        Returns:
            Tuple of (para_idx, translated_runs)
        """
        para_idx, runs, target_language, api_key, model = para_data

        try:
            if runs and any(t.strip() for t in runs):
                translated_runs = self._translate_runs(runs, target_language, api_key, model)
            else:
                translated_runs = runs
            return (para_idx, [(i, t) for i, t in enumerate(translated_runs)])
        except Exception as e:
            print(f"Translation error in paragraph {para_idx}: {str(e)}")
            return (para_idx, [(i, t) for i, t in enumerate(runs)])

    def _translate_paragraph_batch(self, batch_data: tuple) -> dict:
        """
        Translate a batch of paragraphs in parallel.

        Args:
            batch_data: Tuple containing (batch_index, paragraphs_runs_data, target_language, api_key, model, max_workers)

        Returns:
            Dictionary with translated paragraphs
        """
        batch_idx, paragraphs_runs_data, target_language, api_key, model, batch_max_workers = batch_data

        try:
            results = []

            # Prepare data for parallel processing
            para_tasks = [
                (para_idx, runs, target_language, api_key, model)
                for para_idx, runs in paragraphs_runs_data
            ]

            # Process paragraphs in parallel within this batch
            with ThreadPoolExecutor(max_workers=batch_max_workers) as executor:
                results = list(executor.map(self._translate_single_paragraph, para_tasks))

            return {"batch_idx": batch_idx, "translated_paragraphs": results}

        except Exception as e:
            print(f"Translation error in batch {batch_idx}: {str(e)}")
            return {"batch_idx": batch_idx, "error": str(e)}

    def _translate_single_table_paragraph(self, table_para_data: tuple) -> tuple:
        """
        Translate a single table paragraph (to be called in parallel).

        Args:
            table_para_data: Tuple containing (key, runs, target_language, api_key, model)

        Returns:
            Tuple of (key, translated_runs)
        """
        key, runs, target_language, api_key, model = table_para_data

        try:
            if runs and any(t.strip() for t in runs):
                translated_runs = self._translate_runs(runs, target_language, api_key, model)
            else:
                translated_runs = runs
            return (key, [(i, t) for i, t in enumerate(translated_runs)])
        except Exception as e:
            print(f"Translation error in table paragraph {key}: {str(e)}")
            return (key, [(i, t) for i, t in enumerate(runs)])

    def _translate_table_batch(self, batch_data: tuple) -> dict:
        """
        Translate a batch of table paragraphs in parallel.

        Args:
            batch_data: Tuple containing (batch_index, table_para_runs_data, target_language, api_key, model, max_workers)

        Returns:
            Dictionary with translated table paragraphs
        """
        batch_idx, table_para_runs_data, target_language, api_key, model, batch_max_workers = batch_data

        try:
            results = []

            # Prepare data for parallel processing
            table_para_tasks = [
                (key, runs, target_language, api_key, model)
                for key, runs in table_para_runs_data
            ]

            # Process table paragraphs in parallel within this batch
            with ThreadPoolExecutor(max_workers=batch_max_workers) as executor:
                results = list(executor.map(self._translate_single_table_paragraph, table_para_tasks))

            return {"batch_idx": batch_idx, "translated_table_paragraphs": results}

        except Exception as e:
            print(f"Translation error in table batch {batch_idx}: {str(e)}")
            return {"batch_idx": batch_idx, "error": str(e)}

    def translate(self, input_path: str, output_path: str,
                 target_language: str) -> bool:
        """
        Translate Word document using single-level parallelization.

        Args:
            input_path: Path to input DOCX file
            output_path: Path to save translated DOCX file
            target_language: Target language for translation

        Returns:
            Success status
        """
        try:
            doc = Document(input_path)

            # Use API key from the translation client
            api_key = self.client.api_key if hasattr(self.client, 'api_key') and self.client.api_key else os.environ.get('GEMINI_API_KEY', None)

            # Gather ALL paragraphs (body + table) into one flat list with metadata
            all_paragraph_tasks = []

            # Add body paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                runs_texts = [r.text for r in paragraph.runs]
                all_paragraph_tasks.append({
                    'type': 'body',
                    'key': para_idx,
                    'runs': runs_texts,
                    'target_language': target_language,
                    'api_key': api_key,
                    'model': self.client.model
                })

            # Add table cell paragraphs
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            runs_texts = [r.text for r in paragraph.runs]
                            all_paragraph_tasks.append({
                                'type': 'table',
                                'key': (table_idx, row_idx, cell_idx, para_idx),
                                'runs': runs_texts,
                                'target_language': target_language,
                                'api_key': api_key,
                                'model': self.client.model
                            })

            total_paragraphs = len(all_paragraph_tasks)
            max_workers = self.client.max_workers  # 256

            print(f"Processing DOCX with single-level parallelization:")
            print(f"  - Total paragraphs (body + table): {total_paragraphs}")
            print(f"  - All paragraphs processed in parallel with {max_workers} workers")
            print(f"  - Expected concurrent API calls: {min(max_workers, total_paragraphs)}")

            # Process ALL paragraphs in parallel with single worker pool
            translated_results = {}
            if all_paragraph_tasks:
                print(f"Translating {total_paragraphs} paragraphs in parallel...")

                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    # Submit all tasks
                    future_to_task = {
                        executor.submit(
                            self._translate_runs,
                            task['runs'],
                            task['target_language'],
                            task['api_key'],
                            task['model']
                        ): task
                        for task in all_paragraph_tasks
                    }

                    # Collect results as they complete
                    for future in concurrent.futures.as_completed(future_to_task):
                        task = future_to_task[future]
                        try:
                            translated_runs = future.result()
                            # Store with metadata for reconstruction
                            translated_results[task['key']] = {
                                'type': task['type'],
                                'runs': [(i, t) for i, t in enumerate(translated_runs)]
                            }
                        except Exception as e:
                            print(f"Translation failed for {task['type']} paragraph {task['key']}: {str(e)}")
                            # Fallback to original
                            translated_results[task['key']] = {
                                'type': task['type'],
                                'runs': [(i, t) for i, t in enumerate(task['runs'])]
                            }

            # Debug: show keys in translated_results
            print(f"[DEBUG] translated_results has {len(translated_results)} entries")
            body_keys = [k for k in translated_results.keys() if isinstance(k, int)]
            table_keys = [k for k in translated_results.keys() if isinstance(k, tuple)]
            print(f"[DEBUG] Body paragraph keys: {body_keys[:5]}...")  # First 5
            print(f"[DEBUG] Table paragraph keys: {table_keys[:5]}...")  # First 5

            # Apply body paragraph translations
            body_applied = 0
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if para_idx in translated_results:
                    result = translated_results[para_idx]
                    for run_idx, run_text in result['runs']:
                        if run_idx < len(paragraph.runs):
                            old_text = paragraph.runs[run_idx].text
                            paragraph.runs[run_idx].text = run_text
                            body_applied += 1
                            if para_idx < 3:  # Debug first 3 paragraphs
                                print(f"[DEBUG] Applied body para {para_idx}, run {run_idx}: '{old_text}' -> '{run_text}'")
            print(f"[DEBUG] Applied {body_applied} body paragraph runs")

            # Apply table translations
            table_applied = 0
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            key = (table_idx, row_idx, cell_idx, para_idx)
                            if key in translated_results:
                                result = translated_results[key]
                                for run_idx, run_text in result['runs']:
                                    if run_idx < len(paragraph.runs):
                                        paragraph.runs[run_idx].text = run_text
                                        table_applied += 1
            print(f"[DEBUG] Applied {table_applied} table paragraph runs")

            # Save translated document
            doc.save(output_path)
            print(f"[SUCCESS] Translated DOCX saved to: {output_path}")
            return True

        except Exception as e:
            print(f"Error translating DOCX: {str(e)}")
            return False
